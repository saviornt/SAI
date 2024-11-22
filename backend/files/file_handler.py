# backend/files/file_handler.py

import json
import pandas as pd
import pickle
import sqlite3
import shutil
import xml.etree.ElementTree as ET
import cv2
from fpdf import FPDF
from docx import Document
from pydub import AudioSegment
from PyPDF2 import PdfReader
import logging
from .file_type_dictionary import save_operations_dict, load_operations_dict

logger = logging.getLogger(__name__)

class FileHandler:
    def __init__(self):
        self.save_operations = save_operations_dict
        self.load_operations = load_operations_dict

    async def handle_file(self, file_format: str):
        """
        Handle operations based on file type.

        Args:
            file_format (str): The format of the file.

        Returns:
            dict: A dictionary containing save and load handler functions for the specified file format.
        """
        return {
            'save': self.save_operations.get(file_format, self.unsupported_file_format),
            'load': self.load_operations.get(file_format, self.unsupported_file_format)
        }

    async def save_json(self, file_path, data):
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save JSON file '{file_path}': {e}")
            raise

    async def load_json(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (IOError, OSError) as e:
            logger.error(f"Failed to load JSON file '{file_path}': {e}")
            raise

    async def save_text(self, file_path, data):
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(data)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save text file '{file_path}': {e}")
            raise

    async def load_text(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except (IOError, OSError) as e:
            logger.error(f"Failed to load text file '{file_path}': {e}")
            raise

    async def save_csv(self, file_path, data):
        try:
            data.to_csv(file_path, index=False)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save CSV file '{file_path}': {e}")
            raise

    async def load_csv(self, file_path):
        try:
            return pd.read_csv(file_path)
        except (IOError, OSError) as e:
            logger.error(f"Failed to load CSV file '{file_path}': {e}")
            raise

    async def save_xlsx(self, file_path, data):
        try:
            data.to_excel(file_path, index=False)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save Excel file '{file_path}': {e}")
            raise

    async def load_xlsx(self, file_path):
        try:
            return pd.read_excel(file_path)
        except (IOError, OSError) as e:
            logger.error(f"Failed to load Excel file '{file_path}': {e}")
            raise

    async def save_pdf(self, file_path, data):
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, data)
            pdf.output(file_path)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save PDF file '{file_path}': {e}")
            raise

    async def load_pdf(self, file_path):
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except (IOError, OSError) as e:
            logger.error(f"Failed to load PDF file '{file_path}': {e}")
            raise

    async def save_docx(self, file_path, data):
        try:
            doc = Document()
            doc.add_paragraph(data)
            doc.save(file_path)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save DOCX file '{file_path}': {e}")
            raise

    async def load_docx(self, file_path):
        try:
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except (IOError, OSError) as e:
            logger.error(f"Failed to load DOCX file '{file_path}': {e}")
            raise

    async def save_image(self, file_path, data):
        try:
            cv2.imwrite(file_path, data)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save image file '{file_path}': {e}")
            raise

    async def load_image(self, file_path):
        try:
            return cv2.imread(file_path)
        except (IOError, OSError) as e:
            logger.error(f"Failed to load image file '{file_path}': {e}")
            raise

    async def save_mp4(self, file_path, data):
        try:
            with open(file_path, 'wb') as f:
                f.write(data)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save MP4 file '{file_path}': {e}")
            raise

    async def load_mp4(self, file_path):
        try:
            with open(file_path, 'rb') as f:
                return f.read()
        except (IOError, OSError) as e:
            logger.error(f"Failed to load MP4 file '{file_path}': {e}")
            raise

    async def save_audio(self, file_path, data):
        try:
            data.export(file_path, format=file_path.split('.')[-1])
        except (IOError, OSError) as e:
            logger.error(f"Failed to save audio file '{file_path}': {e}")
            raise

    async def load_audio(self, file_path):
        try:
            return AudioSegment.from_file(file_path)
        except (IOError, OSError) as e:
            logger.error(f"Failed to load audio file '{file_path}': {e}")
            raise

    async def save_archive(self, file_path, data):
        try:
            with open(file_path, 'wb') as f:
                f.write(data)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save archive file '{file_path}': {e}")
            raise

    async def load_archive(self, file_path):
        try:
            with open(file_path, 'rb') as f:
                return f.read()
        except (IOError, OSError) as e:
            logger.error(f"Failed to load archive file '{file_path}': {e}")
            raise

    async def save_sqlite(self, file_path, data):
        try:
            conn = sqlite3.connect(file_path)
            cursor = conn.cursor()
            for table_name, df in data.items():
                df.to_sql(table_name, conn, if_exists='replace', index=False)
            conn.commit()
        except (IOError, OSError) as e:
            logger.error(f"Failed to save SQLite database '{file_path}': {e}")
            raise
        finally:
            conn.close()

    async def load_sqlite(self, file_path):
        try:
            conn = sqlite3.connect(file_path)
            cursor = conn.cursor()
            tables = cursor.execute("SELECT name FROM sqlite_master WHERE type='table';").fetchall()
            data = {table[0]: pd.read_sql_query(f"SELECT * FROM {table[0]}", conn) for table in tables}
            return data
        except (IOError, OSError) as e:
            logger.error(f"Failed to load SQLite database '{file_path}': {e}")
            raise
        finally:
            conn.close()

    async def save_pickle(self, file_path, data):
        try:
            with open(file_path, 'wb') as f:
                pickle.dump(data, f)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save pickle file '{file_path}': {e}")
            raise

    async def load_pickle(self, file_path):
        try:
            with open(file_path, 'rb') as f:
                return pickle.load(f)
        except (IOError, OSError) as e:
            logger.error(f"Failed to load pickle file '{file_path}': {e}")
            raise

    async def save_xml(self, file_path, data):
        try:
            tree = ET.ElementTree(data)
            tree.write(file_path)
        except (IOError, OSError) as e:
            logger.error(f"Failed to save XML file '{file_path}': {e}")
            raise

    async def load_xml(self, file_path):
        try:
            tree = ET.parse(file_path)
            return tree.getroot()
        except (IOError, OSError) as e:
            logger.error(f"Failed to load XML file '{file_path}': {e}")
            raise

    async def unsupported_file_format(self, *args, **kwargs):
        logger.error(f"Unsupported file format for operation.")
        raise ValueError("Unsupported file format.")
